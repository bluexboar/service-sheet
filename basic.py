from docx import Document
from docx.shared import Inches
__author__ = 'Will'


page_content_width = 6


if __name__ == '__main__':
    d = Document()
    title = d.add_picture('header.png', width=Inches(page_content_width))
    subtitle = d.add_table(rows=1, cols=2)
    subtitle1 = subtitle.cell(0, 0)
    date = subtitle.cell(0, 1)
    subtitle1.text = '周日下午两点半 ＊＊中文聚会＊＊'
    date.text = '2017年7月30日'
    song1 = d.add_paragraph('第一首歌')

    d.save('basic.docx')