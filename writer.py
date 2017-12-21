__author__ = 'Will'
import shelve
from song_shelve import Song
from docx import Document
from docx.shared import Inches


class Writer:
    def __init__(self, db_name):
        self.db = shelve.open(db_name)
        self.list = []
        self._generate_list()  # generate a list of song titles at the start

    def _generate_list(self):
        # search the database and generate key-title lists
        # sort automatically by key in an increasing order
        key_list = []
        title_list = []
        for key in self.db:
            key_list.append(key)
            title_list.append(self.db[key].title)
        for i in range(len(key_list)):
            pos = key_list.index(str(i))
            self.list.append(title_list[pos])

    def print_list(self, with_limit=True, limit=10):
        # list the titles of songs
        # default list first 10 songs
        for i, item in enumerate(self.list):
            print('[' + str(i) + ']:\t' + item)
            if with_limit:
                limit -= 1
            if limit == 0:
                break

    def _search(self, string):
        # searchable then return song position
        if string not in self.list:
            return False, None
        pos = self.list.index(string)
        return True, pos

    def content(self, string, line_hor_limit=40, line_ver_limit=40):
        # fetch the content of a song, string being the title of the song

        # if the song is searchable:
        search_result, pos = self._search(string)
        if not search_result:
            print('cannot find this song: ' + string)
            return False, None

        # fetch the content
        song = self.db[str(pos)]
        verse = []
        chorus = ''
        verse_line_count = 0  # number of verse lines in total
        chorus_line_count = 0  # number of chorus lines in total
        verse_count = len(song.verse)  # number of verses in total
        # chorus_count = 1  # number of chorus to be repeated
        line_hor_count = []  # character count of the lines

        if song.chorus is not None:
            chorus = song.chorus[0].replace('\nb', '\n\n')
        for i in chorus:
            if i == '\n':
                chorus_line_count += 1
        line_hor_count.append(len(chorus))

        for averse in song.verse:
            averse = averse.replace('\nb', '\n')
            verse.append(averse)
            verse_line_count += 1
            line_hor_count.append(len(averse))
        line_max = max(line_hor_count)  # character count of the longest line

        # generate style
        if line_max >= line_hor_limit:
            print('line too long')
            return False, None

        short_line_count = verse_line_count+chorus_line_count  # minimal line length
        long_line_count = verse_line_count+chorus_line_count * verse_count
        hor_scale = int(line_hor_limit/line_max)
        if hor_scale == 0:
            raise ValueError('line too wide')
        ver_short_scale = int(line_ver_limit/short_line_count)
        ver_long_scale = int(line_ver_limit/long_line_count)

        body = []
        table = dict(col=1, row=1)

        if hor_scale == 1:  # col=1
            if ver_long_scale > 0:  # row > 1
                text = None
                for averse in verse:
                    text += averse + chorus
                body.append(text)
            elif ver_short_scale > 0:
                body.append(verse[0]+chorus+verse[1:])
            else:
                raise ValueError('don\'t fit')
        if hor_scale == 2:
            if verse_count%2 == 0:
                if int(line_ver_limit/long_line_count*2) > 0:
                    table['col'] = 2
                    text1 = None
                    text2 = None
                    for i, averse in enumerate(verse):
                        if i % 2 == 0:
                            text1 += averse+chorus
                        else:
                            text2 += averse+chorus
                    body.append(text1)
                    body.append(text2)
                if int(line_ver_limit/short_line_count*2) > 0:
                    pass

        return True, dict(body=body, table=table, cp=song.copyright)

'''
        d = Document()
        song_table = d.add_table(rows=1, cols=col)
        song_table = song_table.cell(0, 0)
        date = subtitle.cell(0, 1)
        subtitle1.text = '周日下午两点半 ＊＊中文聚会＊＊'
        date.text = '2017年7月30日'
        song1 = d.add_paragraph('第一首歌')


        d.save('basic2.docx')
'''


if __name__ == '__main__':
    w = Writer('song_bank')
    print(w._content('光明之子'))

